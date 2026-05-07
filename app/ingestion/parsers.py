"""Document parsers for PDF, DOCX, and email (.eml / .msg) files.

Each parser returns a list of ``ParsedPage`` objects containing the text
and the original page number so that citations remain accurate.
"""

from __future__ import annotations

import email
import email.policy
import hashlib
import io
from dataclasses import dataclass, field
from pathlib import Path
from typing import Protocol

import pdfplumber
from docx import Document as DocxDocument

from app.core.logging import get_logger

log = get_logger(__name__)


# ── Data contract ─────────────────────────────────────────────────────────────


@dataclass
class ParsedPage:
    """Represents one logical page / section of a document."""

    page_number: int
    text: str
    metadata: dict = field(default_factory=dict)


@dataclass
class ParsedDocument:
    """Full parsing result for one document."""

    pages: list[ParsedPage]
    doc_metadata: dict
    word_count: int
    page_count: int

    @property
    def full_text(self) -> str:
        """Concatenate all pages into a single string."""
        return "\n\n".join(p.text for p in self.pages)

    @property
    def content_hash(self) -> str:
        """SHA-256 hash of the complete extracted text."""
        return hashlib.sha256(self.full_text.encode()).hexdigest()


# ── Parser protocol ───────────────────────────────────────────────────────────


class DocumentParser(Protocol):
    """Protocol that every parser must satisfy."""

    def parse(self, file_bytes: bytes, filename: str) -> ParsedDocument:
        ...


# ── PDF Parser ────────────────────────────────────────────────────────────────


class PDFParser:
    """Extract text and metadata from PDF files using pdfplumber."""

    def parse(self, file_bytes: bytes, filename: str) -> ParsedDocument:
        """Parse PDF bytes into pages.

        Args:
            file_bytes: Raw PDF file content.
            filename: Original filename for logging.

        Returns:
            ParsedDocument with per-page text.
        """
        pages: list[ParsedPage] = []
        doc_metadata: dict = {}

        try:
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                doc_metadata = {
                    "title": pdf.metadata.get("Title", ""),
                    "author": pdf.metadata.get("Author", ""),
                    "creator": pdf.metadata.get("Creator", ""),
                    "creation_date": str(pdf.metadata.get("CreationDate", "")),
                    "page_count": len(pdf.pages),
                }
                for i, page in enumerate(pdf.pages, start=1):
                    text = page.extract_text() or ""
                    text = text.strip()
                    if text:
                        pages.append(ParsedPage(page_number=i, text=text))

            log.info("pdf_parsed", filename=filename, pages=len(pages))
        except Exception as exc:
            log.error("pdf_parse_error", filename=filename, error=str(exc))
            raise RuntimeError(f"Failed to parse PDF '{filename}': {exc}") from exc

        word_count = sum(len(p.text.split()) for p in pages)
        return ParsedDocument(
            pages=pages,
            doc_metadata=doc_metadata,
            word_count=word_count,
            page_count=len(pages),
        )


# ── DOCX Parser ───────────────────────────────────────────────────────────────


class DOCXParser:
    """Extract text from DOCX files using python-docx.

    DOCX files don't have native pages, so we treat every 50 paragraphs
    as a logical "page" to preserve granular citations.
    """

    PARAGRAPHS_PER_PAGE: int = 50

    def parse(self, file_bytes: bytes, filename: str) -> ParsedDocument:
        """Parse DOCX bytes into logical pages.

        Args:
            file_bytes: Raw DOCX content.
            filename: Original filename for logging.

        Returns:
            ParsedDocument with logical-page groupings.
        """
        try:
            doc = DocxDocument(io.BytesIO(file_bytes))
        except Exception as exc:
            log.error("docx_parse_error", filename=filename, error=str(exc))
            raise RuntimeError(f"Failed to parse DOCX '{filename}': {exc}") from exc

        # Extract core properties
        core_props = doc.core_properties
        doc_metadata = {
            "title": core_props.title or "",
            "author": core_props.author or "",
            "subject": core_props.subject or "",
            "created": str(core_props.created or ""),
            "modified": str(core_props.modified or ""),
        }

        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

        # Group into logical pages
        pages: list[ParsedPage] = []
        for page_idx, start in enumerate(
            range(0, max(len(paragraphs), 1), self.PARAGRAPHS_PER_PAGE), start=1
        ):
            chunk_text = "\n".join(paragraphs[start : start + self.PARAGRAPHS_PER_PAGE])
            if chunk_text:
                pages.append(ParsedPage(page_number=page_idx, text=chunk_text))

        if not pages:
            pages.append(ParsedPage(page_number=1, text=""))

        word_count = sum(len(p.text.split()) for p in pages)
        log.info("docx_parsed", filename=filename, pages=len(pages), words=word_count)
        return ParsedDocument(
            pages=pages,
            doc_metadata=doc_metadata,
            word_count=word_count,
            page_count=len(pages),
        )


# ── EML / Email Parser ────────────────────────────────────────────────────────


class EMLParser:
    """Extract text from RFC-2822 .eml email files using Python stdlib."""

    def parse(self, file_bytes: bytes, filename: str) -> ParsedDocument:
        """Parse .eml bytes into a single-page document.

        Args:
            file_bytes: Raw .eml content.
            filename: Original filename for logging.

        Returns:
            ParsedDocument with email body as page 1.
        """
        try:
            msg = email.message_from_bytes(file_bytes, policy=email.policy.default)
        except Exception as exc:
            log.error("eml_parse_error", filename=filename, error=str(exc))
            raise RuntimeError(f"Failed to parse EML '{filename}': {exc}") from exc

        doc_metadata = {
            "from": str(msg.get("From", "")),
            "to": str(msg.get("To", "")),
            "subject": str(msg.get("Subject", "")),
            "date": str(msg.get("Date", "")),
            "message_id": str(msg.get("Message-ID", "")),
        }

        body_parts: list[str] = []
        if msg.is_multipart():
            for part in msg.walk():
                ct = part.get_content_type()
                disp = str(part.get_content_disposition() or "")
                if ct == "text/plain" and "attachment" not in disp:
                    payload = part.get_payload(decode=True)
                    if payload:
                        charset = part.get_content_charset() or "utf-8"
                        body_parts.append(payload.decode(charset, errors="replace"))
        else:
            payload = msg.get_payload(decode=True)
            if payload:
                charset = msg.get_content_charset() or "utf-8"
                body_parts.append(payload.decode(charset, errors="replace"))

        body = "\n".join(body_parts).strip()
        header = (
            f"From: {doc_metadata['from']}\n"
            f"To: {doc_metadata['to']}\n"
            f"Subject: {doc_metadata['subject']}\n"
            f"Date: {doc_metadata['date']}\n\n"
        )
        full_text = header + body

        pages = [ParsedPage(page_number=1, text=full_text)]
        word_count = len(full_text.split())
        log.info("eml_parsed", filename=filename, words=word_count)
        return ParsedDocument(
            pages=pages,
            doc_metadata=doc_metadata,
            word_count=word_count,
            page_count=1,
        )


# ── Parser registry ───────────────────────────────────────────────────────────


_PARSER_MAP: dict[str, DocumentParser] = {
    "pdf": PDFParser(),
    "docx": DOCXParser(),
    "eml": EMLParser(),
    "msg": EMLParser(),  # .msg treated as EML
}


def get_parser(file_extension: str) -> DocumentParser:
    """Return the appropriate parser for *file_extension*.

    Args:
        file_extension: Lowercase file extension without leading dot.

    Raises:
        ValueError: If no parser is registered for the extension.
    """
    ext = file_extension.lower().lstrip(".")
    parser = _PARSER_MAP.get(ext)
    if parser is None:
        raise ValueError(
            f"No parser for extension '{ext}'. "
            f"Supported: {sorted(_PARSER_MAP.keys())}"
        )
    return parser


def parse_document(file_bytes: bytes, filename: str) -> ParsedDocument:
    """Convenience function — auto-detect parser from filename and parse.

    Args:
        file_bytes: Raw file bytes.
        filename: Original filename including extension.

    Returns:
        ParsedDocument ready for chunking.
    """
    ext = Path(filename).suffix.lstrip(".")
    parser = get_parser(ext)
    return parser.parse(file_bytes, filename)
